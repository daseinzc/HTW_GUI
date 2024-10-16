import os
import logging
from docx import Document
from docxcompose.composer import Composer
from docx.enum.text import WD_BREAK

def combine_docx():
    logging.info("合并文档开始")

    input_folder = 'input'
    output_folder = 'output'
    output_file = os.path.join(output_folder, '三联表.docx')

    # 确保输出文件夹存在
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
        logging.info(f"创建输出文件夹: {output_folder}")

    # 读取 input 文件夹中的所有 .docx 文件
    all_file_path = sorted(
        [os.path.join(input_folder, file_name) for file_name in os.listdir(input_folder) if file_name.endswith('.docx')],
        key=lambda x: int(''.join(filter(str.isdigit, os.path.basename(x))))
    )

    if len(all_file_path) == 0:
        logging.error("输入文件夹中没有找到 .docx 文件")
        print("输入文件夹中没有找到 .docx 文件，无法进行合并")
        return

    try:
        # 打开第一个文档作为起始文档
        first_document = Document(all_file_path[0])
        composer = Composer(first_document)

        # 检查 custom.xml 的路径
        current_directory = os.getcwd()
        logging.info(f"当前工作目录: {current_directory}")
        custom_xml_path = os.path.join(current_directory, 'docxcompose/templates/custom.xml')

        if not os.path.exists(custom_xml_path):
            logging.error(f"未找到 custom.xml 文件: {custom_xml_path}")
            print(f"未找到 custom.xml 文件: {custom_xml_path}")
            return
        else:
            logging.info(f"找到 custom.xml 文件: {custom_xml_path}")

        # 合并其他文档并输出进度
        total_files = len(all_file_path)
        for index, file_path in enumerate(all_file_path[1:], start=2):
            try:
                # 添加分页符到文档末尾
                para = first_document.add_paragraph()
                run = para.add_run()
                run.add_break(WD_BREAK.PAGE)

                # 合并文档
                doc = Document(file_path)
                composer.append(doc)
                logging.info(f"成功合并文件: {file_path}")
                print(f"合并进度：{index}/{total_files} 文件 ({file_path})")  # 输出合并进度
            except Exception as e:
                logging.error(f"处理文件 {file_path} 时出错: {e}")

        # 保存合并后的文档到 output 文件夹中
        composer.save(output_file)
        logging.info(f"文档已成功保存到: {output_file}")
        print("文档合并完成，结果已保存到 output 文件夹")

    except Exception as e:
        logging.error(f"合并文档时出错: {e}")
        print("合并文档时出错，请查看日志")

    logging.info("合并文档结束")
