from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.shared import Cm

def single_conduct(school,year,month,money,due_day,i):
    # 各段落文本内容
    p1_word = "财务处："
    p2_word = f'''兹有{school}{year}年{month}月团费{money}元已打至账户：42001227145050000610-0002，特申请转至团委团费经费号：0801613001，请予以支持为荷。'''
    p3_word = "共青团华中科技大学委员会"
    p4_word = "华中科技大学财务处"
    p5_word = "校团委："
    p6_word = school
    p7_word = due_day
    p8_word = f"现已收到{school}{year}年{month}月团费{money}元，特此证明。"
    p9_word = f"现已收到贵单位{year}年{month}月团费{money}元，特此证明。"

    # 创建文档
    doc = Document()

    # 设置文档的默认样式字体
    u = '仿宋_GB2312'
    doc.styles['Normal'].font.size = Pt(16)
    doc.styles['Normal'].font.name = u
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u)

    sections = doc.sections
    for section in sections:
        section.left_margin = Cm(0)  # 将左边距缩小至左侧边缘
        section.right_margin = Cm(3.17)  # 保持默认右边距
        section.top_margin = Cm(2.54)  # 保持默认上边距
        section.bottom_margin = Cm(2.54)  # 保持默认下边距

    # 添加段落
    p1 = doc.add_paragraph(p1_word)
    para_format1 = p1.paragraph_format
    para_format1.left_indent = Cm(3.17)  # 设置文本段落的左缩进到原来左边距的位置
    p2 = doc.add_paragraph(p2_word)
    para_format2 = p2.paragraph_format
    para_format2.left_indent = Cm(3.17)  # 设置文本段落的左缩进到原来左边距的位置
    para_format2.first_line_indent = Pt(16*2)  # 设置段首缩进约为两个字符宽度（以 16pt 字号计算）
    p0 = doc.add_paragraph()
    para_format0 = p0.paragraph_format
    para_format0.left_indent = Cm(3.17)  # 设置文本段落的左缩进到原来左边距的位置
    p3 = doc.add_paragraph(p3_word)
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p4 = doc.add_paragraph(p7_word)
    p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # 插入图片
    p5 = doc.add_picture('fgx.png')
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 继续添加段落
    p6 = doc.add_paragraph(p5_word)
    para_format6 = p6.paragraph_format
    para_format6.left_indent = Cm(3.17)  # 设置文本段落的左缩进到原来左边距的位置
    p7 = doc.add_paragraph(p8_word)
    para_format7 = p7.paragraph_format
    para_format7.left_indent = Cm(3.17)  # 设置文本段落的左缩进到原来左边距的位置
    para_format7.first_line_indent = Pt(16*2)   # 设置段首缩进约为两个字符宽度（以 16pt 字号计算）
    p8 = doc.add_paragraph()
    para_format8 = p8.paragraph_format
    para_format8.left_indent = Cm(3.17)  # 设置文本段落的左缩进到原来左边距的位置
    p9 = doc.add_paragraph(p4_word)
    p9.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p10 = doc.add_paragraph(p7_word)
    p10.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # 插入第二张图片
    p11 = doc.add_picture("fgx.png")
    p11.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 继续添加剩余段落
    p12 = doc.add_paragraph(f"{p6_word}:")
    para_format12 = p12.paragraph_format
    para_format12.left_indent = Cm(3.17)  # 设置文本段落的左缩进到原来左边距的位置
    p13 = doc.add_paragraph(p9_word)
    para_format13 = p13.paragraph_format
    para_format13.left_indent = Cm(3.17)  # 设置文本段落的左缩进到原来左边距的位置
    para_format13.first_line_indent = Pt(16*2)   # 设置段首缩进约为两个字符宽度（以 16pt 字号计算）
    p15 = doc.add_paragraph(p3_word)
    p15.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p16 = doc.add_paragraph(p7_word)
    p16.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # 保存文档
    doc.save(f'./input/{i}{p6_word}.docx')
    return